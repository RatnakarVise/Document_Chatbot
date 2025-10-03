import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
import zipfile
import io

def extract_text_from_pdf(file_obj):
    pdf = PdfReader(file_obj)
    text = ""
    for page in pdf.pages:
        text += page.extract_text() or ""
    return text

def extract_text_from_docx(file_obj):
    doc = Document(file_obj)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

# def extract_text_from_excel(file_obj):
#     df = pd.read_excel(file_obj)
#     text = ""
#     for _, row in df.iterrows():
#         row_text = " | ".join(f"{col}: {row[col]}" for col in df.columns)
#         text += row_text + "\n"
#     return text
def extract_text_from_excel(file_obj, return_df=False):
    df = pd.read_excel(file_obj)
    text = ""
    for i, row in df.iterrows():
        row_text = ". ".join(f"{col}: {row[col]}" for col in df.columns)
        text += f"Row {i+1}: {row_text}\n"
    if return_df:
        return text, df
    return text

def extract_text_from_zip(file_obj):
    text = ""
    with zipfile.ZipFile(file_obj) as z:
        for filename in z.namelist():
            with z.open(filename) as f:
                if filename.endswith(".pdf"):
                    text += extract_text_from_pdf(f)
                elif filename.endswith(".docx"):
                    text += extract_text_from_docx(f)
                elif filename.endswith((".xlsx", ".xls")):
                    file_bytes = f.read()
                    text += extract_text_from_excel(io.BytesIO(file_bytes))
    return text
def get_raw_text(file_bytes, filename):
    raw_text = ""

    filename = filename.lower()

    if filename.endswith(".pdf"):
        from PyPDF2 import PdfReader
        pdf = PdfReader(io.BytesIO(file_bytes))
        for page in pdf.pages:
            raw_text += page.extract_text() or ""

    elif filename.endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            raw_text += para.text + "\n"

    elif filename.endswith((".xlsx", ".xls")):
        import pandas as pd
        df = pd.read_excel(io.BytesIO(file_bytes))
        for i, row in df.iterrows():
            row_text = " | ".join(f"{col}: {row[col]}" for col in df.columns)
            raw_text += row_text + "\n"

    elif filename.endswith(".zip"):
        import zipfile
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            for inner_file in z.namelist():
                with z.open(inner_file) as f:
                    inner_bytes = f.read()
                    raw_text += get_raw_text(inner_bytes, inner_file)  # recursive call

    return raw_text
